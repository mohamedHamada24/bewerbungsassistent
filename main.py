#!/usr/bin/env python3
"""
Bewerbungsassistent
-------------------
Liest einen Lebenslauf als PDF, nimmt eine Stellenbeschreibung entgegen
und erzeugt mit Claude ein deutsches Anschreiben sowie einen angepassten
Lebenslauf – jeweils als DOCX-Datei.

Verwendung:
    python main.py --cv mein_lebenslauf.pdf
    python main.py --cv mein_lebenslauf.pdf --out ./ausgabe
"""

import sys
import argparse
from pathlib import Path
from datetime import date

import anthropic
import pdfplumber
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---------------------------------------------------------------------------
# PDF-Extraktion
# ---------------------------------------------------------------------------

def extract_cv_from_pdf(pdf_path: str) -> str:
    """Extrahiert den Textinhalt aus einer PDF-Lebenslauf-Datei."""
    pages = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
    return "\n\n".join(pages)


# ---------------------------------------------------------------------------
# Eingabe: Stellenbeschreibung
# ---------------------------------------------------------------------------

def read_job_posting() -> str:
    """Fordert den Benutzer auf, eine Stellenbeschreibung einzufügen."""
    print("\n" + "=" * 60)
    print("Stellenbeschreibung einfügen")
    print("(Beenden mit zwei aufeinanderfolgenden leeren Zeilen)")
    print("=" * 60 + "\n")

    lines: list[str] = []
    consecutive_empty = 0

    while True:
        try:
            line = input()
        except EOFError:
            break

        if line == "":
            consecutive_empty += 1
            if consecutive_empty >= 2:
                break
            lines.append("")
        else:
            consecutive_empty = 0
            lines.append(line)

    return "\n".join(lines).strip()


# ---------------------------------------------------------------------------
# Claude: Anschreiben generieren
# ---------------------------------------------------------------------------

ANSCHREIBEN_SYSTEM = """\
Du bist ein erfahrener Karriereberater, spezialisiert auf den deutschsprachigen \
Arbeitsmarkt. Du schreibst professionelle, überzeugende Anschreiben auf Deutsch.

Regeln:
- Kein Absender, kein Empfänger, kein Datum (werden separat hinzugefügt).
- Beginne mit der Betreffzeile (z. B. „Bewerbung als …").
- Klassischer deutscher Briefstil, maximal eine DIN-A4-Seite.
- Konkrete Bezüge zur Stellenbeschreibung und zu den eigenen Qualifikationen.
- Keine Floskeln wie „hiermit bewerbe ich mich".
- Schließe mit „Mit freundlichen Grüßen" ab.
- Gib NUR den Anschreiben-Text aus, keine Erklärungen oder Kommentare.\
"""

def generate_anschreiben(cv_text: str, job_posting: str, client: anthropic.Anthropic) -> str:
    """Lässt Claude ein deutsches Anschreiben generieren."""
    print("\n✍️  Erstelle Anschreiben …")

    with client.messages.stream(
        model="claude-opus-4-7",
        max_tokens=2048,
        system=[{"type": "text", "text": ANSCHREIBEN_SYSTEM, "cache_control": {"type": "ephemeral"}}],
        messages=[
            {
                "role": "user",
                "content": (
                    f"## Mein Lebenslauf:\n{cv_text}\n\n"
                    f"## Stellenbeschreibung:\n{job_posting}\n\n"
                    "Schreibe ein professionelles Anschreiben für diese Bewerbung."
                ),
            }
        ],
    ) as stream:
        message = stream.get_final_message()

    return message.content[0].text


# ---------------------------------------------------------------------------
# Claude: Lebenslauf anpassen
# ---------------------------------------------------------------------------

CV_SYSTEM = """\
Du bist ein Karriereberater und überarbeitest Lebensläufe für spezifische \
Stellenbewerbungen. Du passt den vorhandenen Lebenslauf an – OHNE Fakten zu erfinden.

Anpassungsregeln:
- Behalte ALLE echten Angaben (Daten, Arbeitgeber, Abschlüsse).
- Priorisiere für diese Stelle relevante Erfahrungen.
- Passe Formulierungen an Schlüsselbegriffe der Stellenbeschreibung an.
- Hebe relevante Fähigkeiten hervor.
- Passe einen Profil-/Zusammenfassungsabschnitt an (falls vorhanden).

Ausgabeformat (Markdown – wird in DOCX konvertiert):
- `# Name` für den vollständigen Namen in Zeile 1
- Kontaktdaten in Zeile 2 (E-Mail | Telefon | Ort)
- `## Abschnittsname` für Hauptabschnitte (z. B. ## Berufserfahrung)
- `**Jobtitel** | Unternehmen | Zeitraum` für einzelne Stationen
- `- Aufgabe/Leistung` für Bullet-Points
- Gib NUR den Lebenslauf-Text aus, keine Erklärungen.\
"""

def generate_tailored_cv(cv_text: str, job_posting: str, client: anthropic.Anthropic) -> str:
    """Lässt Claude eine auf die Stelle zugeschnittene Version des Lebenslaufs erstellen."""
    print("📋 Passe Lebenslauf an …")

    with client.messages.stream(
        model="claude-opus-4-7",
        max_tokens=3000,
        system=[{"type": "text", "text": CV_SYSTEM, "cache_control": {"type": "ephemeral"}}],
        messages=[
            {
                "role": "user",
                "content": (
                    f"## Mein aktueller Lebenslauf:\n{cv_text}\n\n"
                    f"## Stelle, auf die ich mich bewerbe:\n{job_posting}\n\n"
                    "Erstelle eine angepasste Version meines Lebenslaufs im Markdown-Format."
                ),
            }
        ],
    ) as stream:
        message = stream.get_final_message()

    return message.content[0].text


# ---------------------------------------------------------------------------
# DOCX-Erstellung: Anschreiben
# ---------------------------------------------------------------------------

def _set_margins(doc: Document, top=2.5, bottom=2.0, left=2.5, right=2.0):
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


def save_anschreiben(text: str, output_path: Path):
    """Speichert das Anschreiben als formatierte DOCX-Datei."""
    doc = Document()
    _set_margins(doc)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    first_content_line = True

    for line in text.split("\n"):
        stripped = line.strip()
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = Pt(14)

        if not stripped:
            continue

        run = para.add_run(stripped)
        run.font.name = "Arial"
        run.font.size = Pt(11)

        # Betreffzeile fett
        if first_content_line or stripped.lower().startswith("betreff"):
            run.bold = True
            first_content_line = False
        else:
            first_content_line = False

    doc.save(str(output_path))
    print(f"   ✅ Anschreiben: {output_path.name}")


# ---------------------------------------------------------------------------
# DOCX-Erstellung: Lebenslauf
# ---------------------------------------------------------------------------

def _add_formatted_text(para, text: str, base_size: Pt = Pt(10.5)):
    """Fügt Text mit **fett**-Unterstützung zu einem Absatz hinzu."""
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)
        run.font.name = "Arial"
        run.font.size = base_size
        run.bold = (i % 2 == 1)  # ungerade Indizes liegen zwischen **


def save_cv(markdown_text: str, output_path: Path):
    """Konvertiert den Markdown-Lebenslauf in eine formatierte DOCX-Datei."""
    doc = Document()
    _set_margins(doc, top=2.0, bottom=2.0, left=2.5, right=2.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    SECTION_COLOR = RGBColor(0x2E, 0x74, 0xB5)  # Blau

    for line in markdown_text.split("\n"):
        stripped = line.strip()

        if not stripped:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            continue

        # H1 – Name
        if stripped.startswith("# "):
            name = stripped[2:].strip()
            para = doc.add_paragraph()
            run = para.add_run(name)
            run.font.name = "Arial"
            run.font.size = Pt(22)
            run.bold = True
            para.paragraph_format.space_after = Pt(2)

        # H2 – Abschnittsüberschriften
        elif stripped.startswith("## "):
            heading = stripped[3:].strip().upper()
            para = doc.add_paragraph()
            run = para.add_run(heading)
            run.font.name = "Arial"
            run.font.size = Pt(11)
            run.bold = True
            run.font.color.rgb = SECTION_COLOR
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(4)
            # Trennlinie unter der Überschrift via Absatz-Rahmen
            from docx.oxml.ns import qn
            from lxml import etree
            pPr = para._p.get_or_add_pPr()
            pBdr = etree.SubElement(pPr, qn("w:pBdr"))
            bottom = etree.SubElement(pBdr, qn("w:bottom"))
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "2E74B5")

        # Bullet-Point
        elif stripped.startswith("- "):
            content = stripped[2:].strip()
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(0.5)
            para.paragraph_format.space_after = Pt(2)
            # Manueller Bullet-Punkt
            run = para.add_run("• ")
            run.font.name = "Arial"
            run.font.size = Pt(10.5)
            _add_formatted_text(para, content)

        # Normaler Text
        else:
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(3)
            _add_formatted_text(para, stripped)

    doc.save(str(output_path))
    print(f"   ✅ Lebenslauf:  {output_path.name}")


# ---------------------------------------------------------------------------
# Einstiegspunkt
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Bewerbungsassistent – Anschreiben & Lebenslauf automatisch erstellen"
    )
    parser.add_argument(
        "--cv", required=True, metavar="LEBENSLAUF.PDF",
        help="Pfad zur PDF-Lebenslauf-Datei"
    )
    parser.add_argument(
        "--out", default=".", metavar="VERZEICHNIS",
        help="Ausgabeverzeichnis (Standard: aktuelles Verzeichnis)"
    )
    args = parser.parse_args()

    cv_path = Path(args.cv)
    if not cv_path.is_file():
        sys.exit(f"❌ Datei nicht gefunden: {args.cv}")
    if cv_path.suffix.lower() != ".pdf":
        sys.exit("❌ Bitte eine PDF-Datei angeben.")

    output_dir = Path(args.out)
    output_dir.mkdir(parents=True, exist_ok=True)

    # --- Lebenslauf einlesen ---
    print(f"\n📄 Lese Lebenslauf: {cv_path.name}")
    cv_text = extract_cv_from_pdf(str(cv_path))
    if not cv_text.strip():
        sys.exit(
            "❌ Kein Text extrahierbar. Handelt es sich um ein gescanntes Bild-PDF?\n"
            "   Tipp: Nutze ein OCR-Tool wie Adobe Acrobat, um das PDF durchsuchbar zu machen."
        )
    print(f"   {len(cv_text):,} Zeichen gelesen.")

    # --- Stellenbeschreibung ---
    job_posting = read_job_posting()
    if not job_posting:
        sys.exit("❌ Keine Stellenbeschreibung eingegeben.")

    # --- Claude generiert die Dokumente ---
    client = anthropic.Anthropic()
    anschreiben_text = generate_anschreiben(cv_text, job_posting, client)
    cv_markdown = generate_tailored_cv(cv_text, job_posting, client)

    # --- Dokumente speichern ---
    today = date.today().strftime("%Y-%m-%d")
    anschreiben_path = output_dir / f"Anschreiben_{today}.docx"
    cv_out_path = output_dir / f"Lebenslauf_angepasst_{today}.docx"

    print("\n💾 Speichere Dokumente …")
    save_anschreiben(anschreiben_text, anschreiben_path)
    save_cv(cv_markdown, cv_out_path)

    print(f"\n🎉 Fertig! Dokumente gespeichert in: {output_dir.resolve()}\n")


if __name__ == "__main__":
    main()
