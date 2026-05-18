"""
Bewerbungsassistent – Streamlit UI mit automatischer Job-Suche
"""

import io
import json
import requests
from datetime import date
from pathlib import Path

CONFIG_FILE = Path(__file__).parent / ".config.json"

def load_config() -> dict:
    if CONFIG_FILE.exists():
        try:
            return json.loads(CONFIG_FILE.read_text())
        except Exception:
            return {}
    return {}

def save_config(data: dict):
    CONFIG_FILE.write_text(json.dumps(data, indent=2))

import anthropic
import pdfplumber
import streamlit as st
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from lxml import etree

# ---------------------------------------------------------------------------
# Seitenkonfiguration
# ---------------------------------------------------------------------------

st.set_page_config(
    page_title="Bewerbungsassistent",
    page_icon="📝",
    layout="wide",
)

# ---------------------------------------------------------------------------
# PDF-Extraktion
# ---------------------------------------------------------------------------

def extract_cv_from_pdf(uploaded_file) -> str:
    pages = []
    with pdfplumber.open(uploaded_file) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
    return "\n\n".join(pages)


# ---------------------------------------------------------------------------
# Job-Suche (LinkedIn)
# ---------------------------------------------------------------------------

@st.cache_resource(show_spinner=False)
def get_linkedin_client(li_at: str, jsessionid: str):
    from linkedin_api import Linkedin
    jar = requests.cookies.RequestsCookieJar()
    jar.set("li_at", li_at, domain=".linkedin.com", path="/")
    jar.set("JSESSIONID", f'"{jsessionid}"' if not jsessionid.startswith('"') else jsessionid,
            domain=".linkedin.com", path="/")
    return Linkedin("", "", cookies=jar)


def search_jobs(query: str, location: str, li_at: str, jsessionid: str, size: int = 20) -> list:
    api = get_linkedin_client(li_at, jsessionid)
    raw = api.search_jobs(
        keywords=query,
        location_name=location,
        limit=size,
    )
    jobs = []
    for item in raw:
        entity = item.get("jobPosting", item)
        job_id = (
            item.get("jobPostingId")
            or item.get("entityUrn", "").split(":")[-1]
        )
        title = (
            entity.get("title")
            or item.get("title", "")
        )
        company = (
            entity.get("companyDetails", {})
            .get("com.linkedin.voyager.deco.jobs.web.shared.WebCompactJobPostingCompany", {})
            .get("companyResolutionResult", {})
            .get("name", "")
            or item.get("companyName", "")
        )
        description = ""
        if job_id:
            try:
                details = api.get_job(job_id)
                description = (
                    details.get("description", {}).get("text", "")
                    or details.get("description", "")
                )
            except Exception:
                pass

        jobs.append({
            "job_title": title,
            "employer_name": company,
            "job_city": location,
            "job_country": "Deutschland",
            "job_description": description or title,
            "job_apply_link": f"https://www.linkedin.com/jobs/view/{job_id}" if job_id else "",
        })
    return jobs


# ---------------------------------------------------------------------------
# Claude – Job-Bewertung
# ---------------------------------------------------------------------------

SCORE_SYSTEM = """\
Du bist ein Karriereberater. Bewerte wie gut eine Stelle zum Lebenslauf passt.
Antworte NUR mit validem JSON – kein Text davor oder danach:
{
  "score": <ganze Zahl 1-10>,
  "gruende": ["Grund 1", "Grund 2", "Grund 3"],
  "zusammenfassung": "Ein Satz warum diese Stelle (nicht) passt."
}
"""

def score_job(cv_text: str, job: dict, client: anthropic.Anthropic) -> dict:
    job_text = (
        f"Stellentitel: {job.get('job_title', '')}\n"
        f"Unternehmen: {job.get('employer_name', '')}\n"
        f"Ort: {job.get('job_city', '')} {job.get('job_country', '')}\n"
        f"Beschreibung:\n{job.get('job_description', '')[:3000]}"
    )
    msg = client.messages.create(
        model="claude-haiku-4-5",
        max_tokens=512,
        system=SCORE_SYSTEM,
        messages=[{
            "role": "user",
            "content": (
                f"## Mein Lebenslauf:\n{cv_text[:3000]}\n\n"
                f"## Stelle:\n{job_text}\n\n"
                "Bewerte die Übereinstimmung."
            ),
        }],
    )
    raw = msg.content[0].text.strip()
    # Strip markdown code fences if present
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
        raw = raw.strip()
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        return {"score": 0, "gruende": [], "zusammenfassung": "Bewertung fehlgeschlagen."}


# ---------------------------------------------------------------------------
# Claude – Streaming-Generatoren
# ---------------------------------------------------------------------------

ANSCHREIBEN_SYSTEM = """\
Du bist ein erfahrener Karriereberater, spezialisiert auf den deutschsprachigen \
Arbeitsmarkt. Du schreibst professionelle, überzeugende Anschreiben auf Deutsch.

Regeln:
- Kein Absender, kein Empfänger, kein Datum (werden separat hinzugefügt).
- Beginne mit der Betreffzeile (z. B. „Bewerbung als …").
- Klassischer deutscher Briefstil, maximal eine DIN-A4-Seite.
- Konkrete Bezüge zur Stellenbeschreibung und zu den eigenen Qualifikationen.
- Keine Floskeln wie „hiermit bewerbe ich mich".
- Schließe mit „Mit freundlichen Grüßen" ab.
- Gib NUR den Anschreiben-Text aus, keine Erklärungen oder Kommentare.\
"""

CV_SYSTEM = """\
Du bist ein Karriereberater und überarbeitest Lebensläufe für spezifische \
Stellenbewerbungen. Du passt den vorhandenen Lebenslauf an – OHNE Fakten zu erfinden.

Anpassungsregeln:
- Behalte ALLE echten Angaben (Daten, Arbeitgeber, Abschlüsse).
- Priorisiere für diese Stelle relevante Erfahrungen.
- Passe Formulierungen an Schlüsselbegriffe der Stellenbeschreibung an.
- Hebe relevante Fähigkeiten hervor.
- Passe einen Profil-/Zusammenfassungsabschnitt an (falls vorhanden).

Ausgabeformat (Markdown):
- `# Name` für den vollständigen Namen in Zeile 1
- Kontaktdaten in Zeile 2 (E-Mail | Telefon | Ort)
- `## Abschnittsname` für Hauptabschnitte
- `**Jobtitel** | Unternehmen | Zeitraum` für Stationen
- `- Aufgabe` für Bullet-Points
- Gib NUR den Lebenslauf-Text aus, keine Erklärungen.\
"""


def stream_anschreiben(cv_text: str, job_posting: str, client: anthropic.Anthropic):
    with client.messages.stream(
        model="claude-haiku-4-5",
        max_tokens=2048,
        system=[{"type": "text", "text": ANSCHREIBEN_SYSTEM, "cache_control": {"type": "ephemeral"}}],
        messages=[{
            "role": "user",
            "content": (
                f"## Mein Lebenslauf:\n{cv_text}\n\n"
                f"## Stellenbeschreibung:\n{job_posting}\n\n"
                "Schreibe ein professionelles Anschreiben für diese Bewerbung."
            ),
        }],
    ) as stream:
        yield from stream.text_stream


def stream_cv(cv_text: str, job_posting: str, client: anthropic.Anthropic):
    with client.messages.stream(
        model="claude-haiku-4-5",
        max_tokens=3000,
        system=[{"type": "text", "text": CV_SYSTEM, "cache_control": {"type": "ephemeral"}}],
        messages=[{
            "role": "user",
            "content": (
                f"## Mein aktueller Lebenslauf:\n{cv_text}\n\n"
                f"## Stelle, auf die ich mich bewerbe:\n{job_posting}\n\n"
                "Erstelle eine angepasste Version meines Lebenslaufs im Markdown-Format."
            ),
        }],
    ) as stream:
        yield from stream.text_stream


# ---------------------------------------------------------------------------
# DOCX-Erstellung
# ---------------------------------------------------------------------------

def _set_margins(doc: Document, top=2.5, bottom=2.0, left=2.5, right=2.0):
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


def build_anschreiben_docx(text: str) -> bytes:
    doc = Document()
    _set_margins(doc)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    first_line = True
    for line in text.split("\n"):
        stripped = line.strip()
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = Pt(14)
        if not stripped:
            continue
        run = para.add_run(stripped)
        run.font.name = "Arial"
        run.font.size = Pt(11)
        if first_line or stripped.lower().startswith("betreff"):
            run.bold = True
        first_line = False

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_formatted_text(para, text: str, size: Pt = Pt(10.5)):
    for i, part in enumerate(text.split("**")):
        if not part:
            continue
        run = para.add_run(part)
        run.font.name = "Arial"
        run.font.size = size
        run.bold = (i % 2 == 1)


def build_cv_docx(markdown_text: str) -> bytes:
    doc = Document()
    _set_margins(doc, top=2.0, bottom=2.0)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    BLUE = RGBColor(0x2E, 0x74, 0xB5)

    for line in markdown_text.split("\n"):
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph()
            continue

        if stripped.startswith("# "):
            para = doc.add_paragraph()
            run = para.add_run(stripped[2:].strip())
            run.font.name = "Arial"
            run.font.size = Pt(22)
            run.bold = True
            para.paragraph_format.space_after = Pt(2)

        elif stripped.startswith("## "):
            para = doc.add_paragraph()
            run = para.add_run(stripped[3:].strip().upper())
            run.font.name = "Arial"
            run.font.size = Pt(11)
            run.bold = True
            run.font.color.rgb = BLUE
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(4)
            pPr = para._p.get_or_add_pPr()
            pBdr = etree.SubElement(pPr, qn("w:pBdr"))
            bottom = etree.SubElement(pBdr, qn("w:bottom"))
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "2E74B5")

        elif stripped.startswith("- "):
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(0.5)
            para.paragraph_format.space_after = Pt(2)
            run = para.add_run("• ")
            run.font.name = "Arial"
            run.font.size = Pt(10.5)
            _add_formatted_text(para, stripped[2:].strip())

        else:
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(3)
            _add_formatted_text(para, stripped)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _generate_docs_for_job(job: dict, cv_text: str, client: anthropic.Anthropic, key_suffix: str):
    job_desc = (
        f"{job.get('job_title', '')} bei {job.get('employer_name', '')}\n"
        f"Ort: {job.get('job_city', '')} {job.get('job_country', '')}\n\n"
        f"{job.get('job_description', '')}"
    )

    st.markdown("**✍️ Anschreiben**")
    anschreiben_text = st.write_stream(stream_anschreiben(cv_text, job_desc, client))

    st.markdown("**📋 Angepasster Lebenslauf**")
    cv_markdown = st.write_stream(stream_cv(cv_text, job_desc, client))

    today = date.today().strftime("%Y-%m-%d")
    employer = job.get("employer_name", "Unternehmen").replace(" ", "_")
    col_a, col_b = st.columns(2)
    with col_a:
        st.download_button(
            "⬇️ Anschreiben.docx",
            data=build_anschreiben_docx(anschreiben_text),
            file_name=f"Anschreiben_{employer}_{today}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key=f"dl_a_{key_suffix}",
        )
    with col_b:
        st.download_button(
            "⬇️ Lebenslauf_angepasst.docx",
            data=build_cv_docx(cv_markdown),
            file_name=f"Lebenslauf_{employer}_{today}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key=f"dl_cv_{key_suffix}",
        )


# ---------------------------------------------------------------------------
# UI
# ---------------------------------------------------------------------------

_cfg = load_config()

with st.sidebar:
    st.header("⚙️ Einstellungen")
    anthropic_key = st.text_input(
        "Anthropic API Key", type="password",
        placeholder="sk-ant-...", value=_cfg.get("anthropic_key", "")
    )
    st.divider()
    st.subheader("LinkedIn Cookie")
    linkedin_li_at = st.text_input(
        "li_at Cookie", type="password",
        placeholder="AQE...",
        value=_cfg.get("linkedin_li_at", ""),
    )
    linkedin_jsessionid = st.text_input(
        "JSESSIONID Cookie", type="password",
        placeholder='ajax:...',
        value=_cfg.get("linkedin_jsessionid", ""),
    )
    st.caption(
        "**Wie bekomme ich die Cookies?**\n\n"
        "1. LinkedIn in Chrome öffnen & einloggen\n"
        "2. F12 → Application → Cookies → linkedin.com\n"
        "3. `li_at` und `JSESSIONID` kopieren"
    )
    if st.button("💾 Zugangsdaten speichern", use_container_width=True):
        save_config({
            "anthropic_key": anthropic_key,
            "linkedin_li_at": linkedin_li_at,
            "linkedin_jsessionid": linkedin_jsessionid,
        })
        st.success("Gespeichert!")
    st.caption("Lokal gespeichert in `.config.json`.")
    st.divider()
    st.header("Anleitung")
    st.markdown(
        """
        1. Beide API Keys eingeben
        2. Lebenslauf als PDF hochladen
        3. Jobtitel & Ort eingeben → **Suchen**
        4. Ergebnisse nach Übereinstimmung sortiert
        5. **Unterlagen erstellen** für passende Jobs
        6. DOCX herunterladen & selbst bewerben
        """
    )
    st.divider()
    st.caption("Job-Suche via LinkedIn · Dokumente via Claude Haiku")

st.title("📝 Bewerbungsassistent")
st.caption("Findet passende Stellen und erstellt Anschreiben + Lebenslauf – automatisch.")
st.divider()

# --- CV Upload ---
uploaded_cv = st.file_uploader("Lebenslauf (PDF)", type=["pdf"])
if uploaded_cv:
    if "cv_file_name" not in st.session_state or st.session_state.cv_file_name != uploaded_cv.name:
        with st.spinner("PDF wird gelesen …"):
            st.session_state.cv_text = extract_cv_from_pdf(uploaded_cv)
            st.session_state.cv_file_name = uploaded_cv.name
    if not st.session_state.cv_text.strip():
        st.error("Kein Text im PDF gefunden. Bitte ein durchsuchbares PDF verwenden.")
        st.stop()
    st.success(f"✅ **{uploaded_cv.name}** gelesen — {len(st.session_state.cv_text):,} Zeichen")

st.divider()

# --- Tabs ---
tab_search, tab_manual = st.tabs(["🔍 Job-Suche", "✏️ Manuelle Eingabe"])

# ===== TAB 1: Job-Suche =====
with tab_search:
    col1, col2 = st.columns(2)
    with col1:
        job_query = st.text_input(
            "Jobtitel / Berufsfeld",
            placeholder="z. B. Software Engineer, Data Analyst, Buchhalter",
        )
    with col2:
        job_location = st.text_input(
            "Ort / Region",
            placeholder="z. B. München, Berlin, Deutschland",
        )

    max_jobs = st.slider("Maximale Anzahl zu bewertender Stellen", 5, 100, 10)
    min_score = st.slider("Mindest-Übereinstimmung anzeigen (1–10)", 1, 10, 5)

    search_btn = st.button("🔍 Jobs suchen & bewerten", type="primary", use_container_width=True)

    if search_btn:
        if not anthropic_key:
            st.error("Bitte den Anthropic API Key in der Seitenleiste eingeben.")
            st.stop()
        if not linkedin_li_at or not linkedin_jsessionid:
            st.error("Bitte li_at und JSESSIONID Cookie in der Seitenleiste eingeben.")
            st.stop()
        if "cv_text" not in st.session_state:
            st.error("Bitte zuerst den Lebenslauf hochladen.")
            st.stop()
        if not job_query.strip() or not job_location.strip():
            st.error("Bitte Jobtitel und Ort eingeben.")
            st.stop()

        client = anthropic.Anthropic(api_key=anthropic_key)
        st.session_state.client = client

        with st.spinner("Suche Stellen …"):
            st.caption(f"Suche: _{job_query}_ in _{job_location}_")
            try:
                raw_jobs = search_jobs(job_query, job_location, linkedin_li_at, linkedin_jsessionid, max_jobs)
            except Exception as e:
                st.error(f"Fehler bei der Jobsuche: {e}")
                st.stop()

        if not raw_jobs:
            st.warning("Keine Stellen gefunden. Versuche andere Suchbegriffe.")
            st.stop()

        jobs_to_score = raw_jobs[:max_jobs]
        st.info(f"**{len(raw_jobs)} Stellen gefunden** – bewerte {len(jobs_to_score)} davon …")

        scored = []
        progress = st.progress(0, text="Bewerte Übereinstimmung …")
        score_error = None
        for i, job in enumerate(jobs_to_score):
            try:
                job["_score"] = score_job(st.session_state.cv_text, job, client)
            except Exception as e:
                score_error = str(e)
                job["_score"] = {"score": 0, "gruende": [], "zusammenfassung": "Bewertung fehlgeschlagen."}
            scored.append(job)
            progress.progress((i + 1) / len(jobs_to_score), text=f"Bewertet: {i + 1}/{len(jobs_to_score)}")
        if score_error:
            st.error(f"Fehler beim Bewerten: {score_error}")

        scored.sort(key=lambda j: j["_score"].get("score", 0), reverse=True)
        st.session_state.scored_jobs = scored
        progress.empty()

    # --- Ergebnisse anzeigen ---
    if "scored_jobs" in st.session_state:
        st.divider()
        visible = [j for j in st.session_state.scored_jobs if j["_score"].get("score", 0) >= min_score]
        st.subheader(f"📊 {len(visible)} passende Stellen (Score ≥ {min_score})")

        if not visible:
            st.info("Keine Stellen über dem Mindest-Score. Schieberegler verringern.")

        for i, job in enumerate(visible):
            score = job["_score"].get("score", 0)
            icon = "🟢" if score >= 7 else "🟡" if score >= 5 else "🔴"
            title = job.get("job_title", "Unbekannte Stelle")
            employer = job.get("employer_name", "Unbekanntes Unternehmen")
            city = job.get("job_city", "")

            with st.expander(f"{icon} **{score}/10** — {title} @ {employer}  |  {city}"):
                st.markdown(f"_{job['_score'].get('zusammenfassung', '')}_")

                gruende = job["_score"].get("gruende", [])
                if gruende:
                    for g in gruende:
                        st.markdown(f"- {g}")

                apply_link = job.get("job_apply_link") or job.get("job_google_link", "")
                if apply_link:
                    st.markdown(f"[🔗 Zur Stellenanzeige]({apply_link})")

                with st.expander("Vollständige Stellenbeschreibung"):
                    st.text(job.get("job_description", "")[:3000])

                st.markdown("---")
                if st.button("📄 Anschreiben & Lebenslauf erstellen", key=f"gen_{i}", use_container_width=True):
                    if "client" not in st.session_state:
                        st.session_state.client = anthropic.Anthropic(api_key=anthropic_key)
                    _generate_docs_for_job(job, st.session_state.cv_text, st.session_state.client, str(i))

# ===== TAB 2: Manuelle Eingabe =====
with tab_manual:
    st.markdown("Stellenbeschreibung direkt einfügen (z. B. von einer Webseite kopiert).")
    job_description = st.text_area(
        "Stellenbeschreibung",
        height=280,
        placeholder="Stellenbeschreibung hier einfügen …",
    )
    generate_btn = st.button("🚀 Unterlagen generieren", type="primary", use_container_width=True)

    if generate_btn:
        if not anthropic_key:
            st.error("Bitte Anthropic API Key in der Seitenleiste eingeben.")
            st.stop()
        if "cv_text" not in st.session_state:
            st.error("Bitte zuerst den Lebenslauf hochladen.")
            st.stop()
        if not job_description.strip():
            st.error("Bitte eine Stellenbeschreibung eingeben.")
            st.stop()

        client = anthropic.Anthropic(api_key=anthropic_key)

        st.subheader("✍️ Anschreiben")
        anschreiben_text = st.write_stream(
            stream_anschreiben(st.session_state.cv_text, job_description, client)
        )

        st.subheader("📋 Angepasster Lebenslauf")
        cv_markdown = st.write_stream(
            stream_cv(st.session_state.cv_text, job_description, client)
        )

        st.divider()
        st.subheader("📥 Dokumente herunterladen")
        today = date.today().strftime("%Y-%m-%d")
        col_a, col_b = st.columns(2)
        with col_a:
            st.download_button(
                "⬇️ Anschreiben.docx",
                data=build_anschreiben_docx(anschreiben_text),
                file_name=f"Anschreiben_{today}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        with col_b:
            st.download_button(
                "⬇️ Lebenslauf_angepasst.docx",
                data=build_cv_docx(cv_markdown),
                file_name=f"Lebenslauf_angepasst_{today}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
